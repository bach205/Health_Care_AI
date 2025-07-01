import shutil
import os
import asyncio
from pathlib import Path
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from src.server.model.documentsModel import save_documents_to_database,save_chunks_into_chroma
from datetime import datetime

UPLOAD_FOLDER = Path(__file__).parent.parent.parent.parent / "documents"

async def save_documents(file,user_id):
    try:
        time_str = datetime.now().strftime("%Y%m%d%H%M%S")
        file.filename = time_str + "_" + file.filename
        # Tạo đường dẫn tuyệt đối để lưu file
        file_location = os.path.join(UPLOAD_FOLDER, file.filename)
        
        # Lưu file, wb la method write + binary (lam viec voi file nhi phan)
        with open(file_location, "wb") as buffer:
            #copy file object of file sang file object of buffer, (chung quy lai la ghi noi dung file vao buffer)
            shutil.copyfileobj(file.file, buffer)
        
        #save documents to database
    
        await save_documents_to_database(file,file_location,user_id)


        asyncio.create_task(index_documents(file_location))
    except Exception as e:
        print(e)
        raise Exception(e)
    return file

async def load_docx_to_text(path):

    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return [text]

async def load_docx_to_documents(path: str):
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

    metadata = {
        "source": os.path.basename(path),
    }

    return [LangchainDocument(page_content=text, metadata=metadata)]

async def chunk_texts(texts):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )
    chunks = []
    for text in texts:
        chunks.extend(splitter.split_text(text))
    return chunks

async def chunk_docs(docs):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )
    chunks = []
    chunks.extend(splitter.split_documents(docs))
    return chunks

#load documents and chunk them into chunks and save them into chroma
async def index_documents(path:str):
    docs = await load_docx_to_documents(path)
    chunks = await chunk_docs(docs)
    await save_chunks_into_chroma(chunks)